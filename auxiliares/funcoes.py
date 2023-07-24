import streamlit as st
import os 
import datetime
from auxiliares.connection import *
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from auxiliares.var_globais import *
from docx.shared import Pt
from documentos.ato.compor_arquivo_atos import preambulo_word, dirigentes

#Baixar Arquivo
def baixar_documento(form_gerado):
    st.markdown('<p style="text-align: center; background: lightgreen">Clique no botão abaixo para fazer download do documento</p>', unsafe_allow_html=True)
    pos_btn_1, pos_btn_2, pos_btn_3 = st.columns([3,3,2])
    with pos_btn_1:
        pass
    with pos_btn_2:
       try:
            with open(str(form_gerado), "rb") as file:
                st.download_button(
                    label="BAIXAR DOCUMENTO",
                    data=file,
                    file_name=form_gerado,
                    mime="text/html")
       except:
            st.error("ARQUIVO NÃO LOCALIZADO! REPITA O PROCESSO.")
    with pos_btn_3:
        pass
    os.remove(form_gerado)
   
# função para deixar data recebida no formato necessário
def data_convertida_br(dt): # recebe uma String
  dia = dt[8:]
  mes = dt[5:7]
  ano = dt[0:4]
  if dia == "":
    return ""
  else:
    return f'{dia}/{mes}/{ano}' # retorna uma String

# gerar word
def gerar_conteudo_publicacao(num_ato, ano_ato, texto, dirigente_responsavel, setor_responsavel): 

    # posicionando inicio do texto
    inicio = texto.find('>') + 1 
    # coletando informações para gravar
    data_atual = data_convertida_br(str(datetime.date.today()))
    _numero_ato_ = f'Nº {num_ato}/{ano_ato}'
    texto_gravar = texto[inicio:].replace('<p style="font-size:12pt;font-family:Calibri;text-indent:25mm;text-align:justify;word-wrap:normal;margin:6pt;">', '').replace('</p>', '').lstrip()

    # mostrando texto do ato abaixo do formulário
    texto_print = f"""{_numero_ato_} - {texto_gravar}"""
    texto_publ = {'Texto_para_Publicação': texto_print}
    st.write(texto_publ)
    ato_gravacao = {'data_emissao': data_atual,
                    'num_formatado': _numero_ato_,
                    'texto_do_ato_gravar': texto_gravar,
                    'dirigente_responsavel': dirigente_responsavel,
                    'setor_responsavel': setor_responsavel}
    return ato_gravacao

def gerar_word_publicacao(word_dict):# , preambulo_doc
    # Criando documento
    documento = Document()

    # core_properties = documento.core_properties / core_properties.language = "pt-BR" -> não funciona para alterar o idioma e corrigir a verificação ortográfica
    # Copiei a solução do github (link abaixo)
    # https://github.com/python-openxml/python-docx/issues/727  
    for my_style in documento.styles:
        style = documento.styles[my_style.name]
        rpr = style.element.get_or_add_rPr()
        lang = docx.oxml.shared.OxmlElement('w:lang')
        if not rpr.xpath('w:lang'):
            lang.set(docx.oxml.shared.qn('w:val'),'de-DE')
            lang.set(docx.oxml.shared.qn('w:eastAsia'),'en-US')
            lang.set(docx.oxml.shared.qn('w:bidi'),'ar-SA')
            rpr.append(lang)  

    # criando cabeçalho do documento
    dia_do_mes = data_atual.day
    nome_do_mes = lista_meses_ano.get(data_atual.month)
    ano_data = data_atual.year
    setor = word_dict['SETOR'][0]

    # configurando texto do título do documento
    data_doc_word = f'ATOS DO(A) {setor}, {dia_do_mes} DE {str(nome_do_mes).upper()} DE {ano_data}\n'
    preambulo_doc_word = None
    assinatura = None

    # configurando nome e tamanho da fonte
    style = documento.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # percorrendo a coluna DIRIGENTE para configurar o preambulo e a assinatura
    for posicao in range(len(word_dict['DIRIGENTE'])):
        if word_dict['DIRIGENTE'][posicao] == "Decano(a) titular":
            preambulo_doc_word = preambulo_word['funcao']['txt_dgp']
            assinatura = dirigentes['dgp']['decanato']['titular']
        elif word_dict['DIRIGENTE'][posicao] == "Decano(a) em exercício":
            preambulo_doc_word = preambulo_word['funcao']['txt_dgp_substituto']
            assinatura = dirigentes['dgp']['decanato']['em_exercicio']
        elif word_dict['DIRIGENTE'][posicao] == "Diretor(a) titular":
            preambulo_doc_word = preambulo_word['funcao']['txt_dgp_dap']
            assinatura = dirigentes['dgp']['dap']['titular']
        elif word_dict['DIRIGENTE'][posicao] == "Diretor(a) em exercício":
            preambulo_doc_word = preambulo_word['funcao']['txt_dgp_dap_susbstituto']
            assinatura = dirigentes['dgp']['dap']['em_exercicio']
        elif word_dict['DIRIGENTE'][posicao] == "Reitor(a)":
            preambulo_doc_word = preambulo_word['funcao']['txt_reitoria']
            assinatura = dirigentes['reitoria']['titular']
        elif word_dict['DIRIGENTE'][posicao] == "Vice-Reitor(a)":
            preambulo_doc_word = preambulo_word['funcao']['txt_reitoria_substituto']
            assinatura = dirigentes['reitoria']['em_exercicio']

    # inserindo cabeçalho (data) no documento
    paragrafo = documento.add_paragraph(data_doc_word)
    paragrafo_formatado = paragrafo.paragraph_format
    paragrafo_formatado.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # inserindo cabeçalho (preambulo) no documento
    paragrafo = documento.add_paragraph(preambulo_doc_word)
    paragrafo_formatado = paragrafo.paragraph_format
    paragrafo_formatado.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
  
    # Gravando texto e gerando documento
    for posicao in range(len(word_dict['NUMERO'])):
        textis = f"{word_dict['NUMERO'][posicao]} - {word_dict['TEXTO'][posicao]}".replace('\n', ' ')
        paragrafo = documento.add_paragraph(textis)
        paragrafo_formatado = paragrafo.paragraph_format
        paragrafo_formatado.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # quebra de página para não acumular tudo em um documento só
        if (posicao+1) == 6 or (posicao+1) == 12:
            documento.add_page_break()

    
    # Assinatura
    assinatura = '\n{0}'.format(assinatura.replace('<br>', '\n'))
    paragrafo = documento.add_paragraph(assinatura)
    paragrafo_formatado = paragrafo.paragraph_format
    paragrafo_formatado.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # aplicando estilos (nome da fonte e tamanho da fonte)
    paragrafo.style = documento.styles['Normal']

    # Salvando documento
    documento.save("AtosPublicacao.docx")

    # baixar documento
    baixar_documento("AtosPublicacao.docx")


def btn_gerar_publicacao(dicionario_publicacao):
    col_btn_1, col_btn_2, col_btn_3= st.columns([2,3,2])
    with col_btn_1:
        pass
    with col_btn_2:
        btn_gerar_word = st.button("GERAR 'WORD' PARA PUBLICAÇÃO")
    with col_btn_3:
        pass
    if btn_gerar_word:
        word_dict = dict(dicionario_publicacao.sort_values(by='NUMERO')) 
        gerar_word_publicacao(word_dict)


def gerar_documento_publicacao():
    st.markdown("<br><h6 style='padding: 5px ;text-align: center; border-style: solid; border-color: brown; border-width: 0.5px;'>Gerar documento para publicação</h6>", unsafe_allow_html=True)

    # criando bando de dados, tabela, inserção e busca
    data_publicacao = st.date_input('Informe a data de emissão do ato: ')
    conectar = criar_conexao('atosgerados_db')
    if data_publicacao != "":
        data_publicacao = data_convertida_br(str(data_publicacao))
        dicionario_publicacao = selecionar_atos(conectar,'atos_publicacao', data_publicacao)
        if dicionario_publicacao.empty:
            st.info("Nenhum Ato foi registrado na data selecionada.")
        else:
            setor = st.selectbox("Selecione o Setor Responsável pela emissão do ato:", lista_de_setores)
            dirigente = None
            if setor == 'DGP':
                dicionario_publicacao = dicionario_publicacao.query('SETOR=="{0}"'.format(setor))
                if dicionario_publicacao.empty:
                    st.info("Nenhum Ato encontrado para o setor selecionado.")
                else:    
                    dirigente = st.selectbox("Selecione o Dirigente Responsável pela emissão do ato:", (dict_dirigentes_responsaveis['dgp']['decanato'][dirig] for dirig in dict_dirigentes_responsaveis['dgp']['decanato'])) # 
                    dicionario_publicacao = dicionario_publicacao.query('DIRIGENTE=="{0}"'.format(dirigente))
                    if dicionario_publicacao.empty:
                        st.info("Nenhum Ato encontrado para o dirigente selecionado.")
                    else:
                        st.table(dicionario_publicacao.sort_values(by='NUMERO')) # Mostrando dados após filtros
                        btn_gerar_publicacao(dicionario_publicacao)
            elif setor == 'DGP/DAP':
                dicionario_publicacao = dicionario_publicacao.query('SETOR=="{0}"'.format(setor))
                if dicionario_publicacao.empty:
                    st.info("Nenhum Ato encontrado para o setor selecionado.")
                else:  
                    dirigente = st.selectbox("Selecione o Dirigente Responsável pela emissão do ato:", (dict_dirigentes_responsaveis['dgp']['dap'][dirig] for dirig in dict_dirigentes_responsaveis['dgp']['decanato'])) # 
                    dicionario_publicacao = dicionario_publicacao.query('DIRIGENTE=="{0}"'.format(dirigente))
                    if dicionario_publicacao.empty:
                        st.info("Nenhum Ato encontrado para o dirigente selecionado.")
                    else:
                        st.table(dicionario_publicacao.sort_values(by='NUMERO')) # Mostrando dados após filtros
                        btn_gerar_publicacao(dicionario_publicacao)
            elif setor == 'REITORIA': 
                dicionario_publicacao = dicionario_publicacao.query('SETOR=="{0}"'.format(setor))
                if dicionario_publicacao.empty:
                    st.info("Nenhum Ato encontrado para o setor selecionado.")
                else:   
                    dirigente = st.selectbox("Selecione o Dirigente Responsável pela emissão do ato:", (dict_dirigentes_responsaveis['reitoria'][dirig] for dirig in dict_dirigentes_responsaveis['reitoria'])) # 
                    dicionario_publicacao = dicionario_publicacao.query('DIRIGENTE=="{0}"'.format(dirigente))
                    if dicionario_publicacao.empty:
                        st.info("Nenhum Ato encontrado para o dirigente selecionado.")
                    else:
                        st.table(dicionario_publicacao.sort_values(by='NUMERO')) # Mostrando dados após filtros
                        btn_gerar_publicacao(dicionario_publicacao)
            else:
                dirigente = st.selectbox("Selecione o Dirigente Responsável pela emissão do ato:",["Escolha o setor responsável."], disabled=True)
                # Mostrando todos os dados     
                st.table(dicionario_publicacao.sort_values(by='NUMERO'))
        

